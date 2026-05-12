"""
Generates outputs/memo.docx — the PM-facing scenario update memo.

The prose is hand-authored and stored as constants in this module. No LLM
call. The build step parses inline asterisk markers into italic runs and
writes a Word document with Times New Roman, 1.5 line spacing, and a small
scenarios table.

Style rules enforced by `_scan_banned`:
  * no em dashes
  * no hyphens in compound modifiers (e.g., write "buy side" not "buy-side")
  * no banned phrases (delve, leverage as verb, robust, significant, etc.)

If any banned token is found at build time, the build raises. The prose has
been audited by hand against the same scanner before being committed here.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
)


TITLE: str = "AEIS Scenario Update Following Q1 2026 Earnings"


# ----- Swappable copy variants ----------------------------------------------
# Two opening renderings for the Summary block. OPENING_A is one long sentence
# preserving the original cadence; OPENING_B splits it into three shorter
# sentences. Change SUMMARY_BLOCK below to flip.

OPENING_A: str = (
    "The *$363* probability weighted target sits *6 percent* below current "
    "*$387*, with the *$368* base case essentially at spot, so the stock is "
    "already pricing the bull side of the distribution at our equal weighted "
    "33/33/33 probability framework. At current levels the framework favors "
    "*trimming long exposure* or waiting for either a *DC growth print above "
    "40 percent* or a *multiple reset toward peers* before adding."
)

OPENING_B: str = (
    "The *$363* probability weighted target sits *6 percent* below current "
    "*$387*. The *$368* base case is essentially at spot, so the stock is "
    "already pricing the bull side of the distribution at our 33/33/33 "
    "weighting. Multiple compression swings the call. At current levels the "
    "framework favors *trimming long exposure* or waiting for either a *DC "
    "growth print above 40 percent* or a *multiple reset toward peers* "
    "before adding."
)

SUMMARY_BLOCK: str = OPENING_A


# Three variants for "What Would Change My View." Lead with bull triggers,
# then bear triggers. No "lift us to bull" / "pull us back to bear" parallel
# construction. No semicolons. Change VIEW_BLOCK below to flip.

VIEW_VARIANT_A: str = (
    "Bull case opens up if CY27 DC growth prints above *40 percent* and BAC "
    "consensus moves toward *$13* EPS. A second hyperscaler Kyber 800V "
    "design win in 2H 2026 is the cleanest single catalyst. Bear case takes "
    "hold on Q2 DC "
    "moderation below mid 30 percent, softening Industrial and Medical bookings off the *+14 "
    "percent* QoQ Q1 base, or a BAC PT cut back toward *$400*."
)

VIEW_VARIANT_B: str = (
    "Bullish trigger set is wide. CY27 DC growth above *40 percent*, BAC "
    "consensus toward *$13* EPS, or a second hyperscaler Kyber 800V design "
    "win in 2H 2026 each move us up. The bearish path is narrower and runs "
    "through Q2 "
    "DC moderation below mid 30 percent, with a BAC PT cut to *$400* as "
    "the cleanest confirmation."
)

VIEW_VARIANT_C: str = (
    "Three signals open the bull case: CY27 DC growth above *40 percent*, "
    "BAC consensus EPS toward *$13*, or a second hyperscaler Kyber 800V "
    "design win in 2H 2026. The bear path is narrower. A Q2 DC moderation "
    "below "
    "mid 30 percent, or a BAC PT cut back toward *$400*, breaks the "
    "multiple compression argument."
)

VIEW_BLOCK: str = VIEW_VARIANT_A


# Section prose. Inline *asterisks* mark italic runs. Section titles are
# empty strings for blocks that should appear without a section header.
SECTIONS: list[tuple[str, list[str]]] = [
    ("Summary", [SUMMARY_BLOCK]),
    ("", ["__TABLE__"]),
    (
        "What Changed Since May 1",
        [
            "Q1 cleared at *$511M* revenue and *$2.09* non GAAP EPS, both "
            "above guide midpoint. Q2 guide of *$540M* and *$2.18* EPS came "
            "in ahead of Street, with Datacenter Computing revenue more than "
            "doubling year over year. BAC raised CY27 EPS to *$12.00* from "
            "*$9.90* on stronger leading edge memory and logic exposure, "
            "lifted the multiple to *36x* from *33x*, and took PT to *$430*. "
            "Cowen raised CY27 EPS to *$12.50* from *$10.30* on the same "
            "earnings power but held the multiple at *28x*, lifting PT to "
            "*$350* while staying Hold on valuation. Bloomberg consensus "
            "settles at *$10.83*.",
        ],
    ),
    (
        "The Multiple Debate",
        [
            "AEIS trades at a *25 percent* PE premium to AMAT and LRCX, well "
            "above the *5 percent* three year average that Cowen anchors as "
            "the fair value range, with comps in the *25 to 26x* range "
            "versus AEIS near *32x* at spot. Either DC growth reaccelerates "
            "from mid 30 "
            "percent toward *40 percent*, or that premium compresses back "
            "toward the historical band. The math is binary. At those "
            "scenario prices and a *33 percent* base weight, current *$387* "
            "implies the market is weighting bull near *47 percent* and bear "
            "near *20 percent*, materially more bullish than our equal "
            "weighted prior, which means the stock is paying for upside that "
            "has not yet printed in earnings or backlog.",
        ],
    ),
    ("What Would Change My View", [VIEW_BLOCK]),
]

SCENARIOS_TABLE: list[list[str]] = [
    ["Scenario", "EPS (CY27)", "P/E", "Price", "Prob"],
    ["Bear", "$10.83", "25x", "$271", "33%"],
    ["Base", "$12.25", "30x", "$368", "33%"],
    ["Bull", "$12.50", "36x", "$450", "33%"],
]

BANNED_PHRASES: list[str] = [
    "delve", "leverage", "robust", "significant", "furthermore", "moreover",
    "additionally", "going forward", "navigate", "strategically positioned",
    "in today's", "it's worth noting", "key takeaway", "navigate the landscape",
    "against this backdrop", "broadly speaking", "in essence",
]


# ----- Inline italic parser -------------------------------------------------

def _parse_inline(text: str) -> list[tuple[str, bool]]:
    """Split text on *italic* markers into (chunk, is_italic) tuples."""
    out: list[tuple[str, bool]] = []
    buf: list[str] = []
    italic = False
    for ch in text:
        if ch == "*":
            if buf:
                out.append(("".join(buf), italic))
                buf = []
            italic = not italic
        else:
            buf.append(ch)
    if buf:
        out.append(("".join(buf), italic))
    return out


# ----- Banned-phrase scanner ------------------------------------------------

def _scan_banned(text: str) -> list[str]:
    """Return all banned-phrase findings in `text`. Empty list means clean."""
    findings: list[str] = []
    lower = text.lower()
    for phrase in BANNED_PHRASES:
        pattern = r"\b" + re.escape(phrase.lower()) + r"\b"
        if re.search(pattern, lower):
            findings.append(phrase)
    if "—" in text or "—" in text:
        findings.append("em dash")
    hyphenated = re.findall(r"[A-Za-z]+-[A-Za-z]+", text)
    if hyphenated:
        findings.append(f"hyphenated_compound: {hyphenated}")
    return findings


def _word_count(text: str) -> int:
    """Count whitespace-separated word tokens (asterisks stripped)."""
    plain = text.replace("*", "")
    return len(re.findall(r"\S+", plain))


# ----- Docx builders --------------------------------------------------------

# Page-fit knobs. Bumped down by `write_memo` if the document spills onto a
# second page. Body font stays at 12pt per spec.
LINE_SPACING_LEVELS: list[float] = [1.5, 1.25]
MARGIN_LEVELS_INCHES: list[float] = [1.0, 0.75]


def _set_run(run: "object", *, bold: bool = False, italic: bool = False,
             size_pt: int = 12) -> None:
    """Apply Times New Roman font with the requested style to a docx run."""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def _add_title(doc: DocxDocument, text: str, line_spacing: float) -> None:
    """Add the centered 14pt bold italic title."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    _set_run(run, bold=True, italic=True, size_pt=14)


def _add_section_header(doc: DocxDocument, text: str,
                         line_spacing: float) -> None:
    """Add a left-aligned bold (non-italic) 12pt section header."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    _set_run(run, bold=True, italic=False, size_pt=12)


def _add_body_paragraph(doc: DocxDocument, text: str,
                         line_spacing: float) -> None:
    """Add a body paragraph with inline italic markers parsed into runs."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = line_spacing
    for chunk, italic in _parse_inline(text):
        run = p.add_run(chunk)
        _set_run(run, italic=italic)


def _add_scenarios_table(doc: DocxDocument, rows: list[list[str]]) -> None:
    """Add the bear/base/bull table, horizontally centered on the page."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            _set_run(run, bold=(r_idx == 0), size_pt=11)


# ----- PDF builder (reportlab) ----------------------------------------------

def _inline_to_rl(text: str) -> str:
    """Convert *italic* markers into reportlab's <i></i> inline markup."""
    out: list[str] = []
    italic = False
    for ch in text:
        if ch == "*":
            out.append("</i>" if italic else "<i>")
            italic = not italic
        else:
            out.append(ch)
    return "".join(out)


def _pdf_styles() -> dict:
    """Return the four ParagraphStyle objects used in the memo PDF."""
    return {
        "title": ParagraphStyle(
            "Title", fontName="Times-BoldItalic", fontSize=14,
            alignment=1, leading=14 * 1.5, spaceAfter=10,
        ),
        "section": ParagraphStyle(
            "Section", fontName="Times-Bold", fontSize=12,
            alignment=0, leading=12 * 1.5, spaceBefore=8, spaceAfter=2,
        ),
        "body": ParagraphStyle(
            "Body", fontName="Times-Roman", fontSize=12,
            alignment=0, leading=12 * 1.5, spaceAfter=4,
        ),
    }


def _scenarios_pdf_table() -> Table:
    """Build the bear/base/bull table, horizontally centered on the page."""
    t = Table(SCENARIOS_TABLE, hAlign="CENTER")
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#305496")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BFBFBF")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def _pdf_styles_with(line_spacing: float) -> dict:
    """Return ParagraphStyle objects parameterized by line spacing."""
    return {
        "title": ParagraphStyle(
            "Title", fontName="Times-BoldItalic", fontSize=14,
            alignment=1, leading=14 * line_spacing, spaceAfter=10,
        ),
        "section": ParagraphStyle(
            "Section", fontName="Times-Bold", fontSize=12,
            alignment=0, leading=12 * line_spacing, spaceBefore=8,
            spaceAfter=2,
        ),
        "body": ParagraphStyle(
            "Body", fontName="Times-Roman", fontSize=12,
            alignment=0, leading=12 * line_spacing, spaceAfter=4,
        ),
    }


def write_memo_pdf(out_path: Path, line_spacing: float = 1.5,
                   margin_inches: float = 1.0) -> int:
    """
    Compose outputs/memo.pdf using the same prose constants as the docx.
    Returns the page count so the caller can decide whether to tighten.
    """
    styles = _pdf_styles_with(line_spacing)
    doc = SimpleDocTemplate(
        str(out_path), pagesize=letter,
        leftMargin=margin_inches * inch, rightMargin=margin_inches * inch,
        topMargin=margin_inches * inch, bottomMargin=margin_inches * inch,
        title="AEIS Scenario Update",
    )
    story: list = [Paragraph(TITLE, styles["title"]), Spacer(1, 6)]
    for section_title, blocks in SECTIONS:
        if section_title:
            story.append(Paragraph(section_title, styles["section"]))
        for block in blocks:
            if block == "__TABLE__":
                story.append(Spacer(1, 2))
                story.append(_scenarios_pdf_table())
                story.append(Spacer(1, 4))
            else:
                story.append(Paragraph(_inline_to_rl(block), styles["body"]))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)
    # Re-open to count pages
    import pdfplumber
    with pdfplumber.open(out_path) as pdf:
        return len(pdf.pages)


# ----- Build entry point ----------------------------------------------------

def write_memo(out_path: Path) -> dict:
    """
    Compose outputs/memo.docx (and the sibling .pdf) and return a build report
    including word count and banned-phrase scan results.
    """
    all_prose = "\n".join(
        block for _, blocks in SECTIONS for block in blocks if block != "__TABLE__"
    )
    findings = _scan_banned(all_prose)
    word_count = _word_count(all_prose)

    if findings:
        raise ValueError(
            f"Memo prose failed banned-phrase scan: {findings}. "
            f"Fix the prose constants in src/memo.py."
        )

    # Page-fit fallback: try (line_spacing=1.5, margin=1.0); if PDF spills to
    # 2 pages, tighten line spacing to 1.25; if still spills, drop margins to
    # 0.75 inch. The docx mirrors whichever level was chosen.
    pdf_path = out_path.with_suffix(".pdf")
    chosen_spacing: float = LINE_SPACING_LEVELS[0]
    chosen_margin: float = MARGIN_LEVELS_INCHES[0]
    pages = write_memo_pdf(pdf_path, line_spacing=chosen_spacing,
                           margin_inches=chosen_margin)
    if pages > 1:
        chosen_spacing = LINE_SPACING_LEVELS[1]
        pages = write_memo_pdf(pdf_path, line_spacing=chosen_spacing,
                               margin_inches=chosen_margin)
    if pages > 1:
        chosen_margin = MARGIN_LEVELS_INCHES[1]
        pages = write_memo_pdf(pdf_path, line_spacing=chosen_spacing,
                               margin_inches=chosen_margin)

    # Now build the matching docx with the same line spacing
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(chosen_margin * 72)
        section.bottom_margin = Pt(chosen_margin * 72)
        section.left_margin = Pt(chosen_margin * 72)
        section.right_margin = Pt(chosen_margin * 72)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    _add_title(doc, TITLE, line_spacing=chosen_spacing)
    for section_title, blocks in SECTIONS:
        if section_title:
            _add_section_header(doc, section_title, line_spacing=chosen_spacing)
        for block in blocks:
            if block == "__TABLE__":
                _add_scenarios_table(doc, SCENARIOS_TABLE)
            else:
                _add_body_paragraph(doc, block, line_spacing=chosen_spacing)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    return {
        "path": str(out_path),
        "pdf_path": str(pdf_path),
        "word_count": word_count,
        "banned_findings": findings,
        "sections": [s for s, _ in SECTIONS],
        "line_spacing": chosen_spacing,
        "margin_inches": chosen_margin,
        "pdf_pages": pages,
    }


def memo_plain_text() -> str:
    """Return the memo prose as plain text (asterisks stripped). For preview."""
    lines: list[str] = [TITLE, ""]
    for section_title, blocks in SECTIONS:
        lines.append(section_title)
        for block in blocks:
            if block == "__TABLE__":
                widths = [max(len(row[c]) for row in SCENARIOS_TABLE)
                          for c in range(len(SCENARIOS_TABLE[0]))]
                for row in SCENARIOS_TABLE:
                    lines.append("  " + "  ".join(
                        cell.ljust(w) for cell, w in zip(row, widths)
                    ))
            else:
                lines.append(block.replace("*", ""))
        lines.append("")
    return "\n".join(lines)
